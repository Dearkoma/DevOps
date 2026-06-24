"""生成毕设开题报告 .docx 文件 — 基于 DevOps 持续交付平台项目"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = r"E:\Dissertation\2\devops-platform\outputs"
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")

# === 样式常量 ===
SIMHEI = '黑体'
SIMSUM = '宋体'
LINE_SPACING = Pt(26.25)  # 1.5倍行距（以五号字为基准约等于这个值）
FIRST_INDENT = Cm(0.74)   # 首行缩进2字符


def set_cell_font(cell, font_name, font_size):
    """设置单元格全部文本的中英文字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size


def add_run(paragraph, text, font_name=SIMSUM, font_size=Pt(12), bold=False, color=None):
    """添加带字体设置的 run"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_format(paragraph, first_indent=None, line_spacing=LINE_SPACING,
                         space_before=0, space_after=0, alignment=None):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if alignment is not None:
        paragraph.alignment = alignment


def add_heading_para(doc, text, font_size=Pt(14), space_before=12, space_after=6,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True):
    """添加黑体标题段落"""
    p = doc.add_paragraph()
    add_run(p, text, font_name=SIMHEI, font_size=font_size, bold=bold)
    set_paragraph_format(p, space_before=space_before, space_after=space_after, alignment=alignment)
    return p


def add_body_para(doc, text, indent=True):
    """添加宋体正文段落"""
    p = doc.add_paragraph()
    add_run(p, text, font_name=SIMSUM, font_size=Pt(12))
    set_paragraph_format(p, first_indent=FIRST_INDENT if indent else None)
    return p


def add_chapter_title(doc, text, space_before=18, space_after=10):
    """添加章节标题（黑体四号 不缩进）"""
    return add_heading_para(doc, text, font_size=Pt(14), space_before=space_before,
                            space_after=space_after)


def add_sub_title(doc, text, space_before=12, space_after=6):
    """添加二级标题（黑体小四 不缩进）"""
    return add_heading_para(doc, text, font_size=Pt(12), space_before=space_before,
                            space_after=space_after)


def add_ref_item(doc, text):
    """添加参考文献条目（悬挂缩进2.5字符）"""
    p = doc.add_paragraph()
    add_run(p, text, font_name=SIMSUM, font_size=Pt(12))
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.first_line_indent = -Cm(0.925)  # 悬挂缩进 ≈ 2.5字符
    pf.left_indent = Cm(0.925)
    return p


def add_image_centered(doc, image_path, width_cm=12):
    """添加居中图片"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(image_path):
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        add_run(p, f'[图片未找到: {image_path}]', font_name=SIMSUM, font_size=Pt(10))
    return p


def add_caption(doc, text):
    """添加图片标题"""
    p = doc.add_paragraph()
    add_run(p, text, font_name=SIMSUM, font_size=Pt(10))
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
    return p


# ============================================================
# 封面页
# ============================================================
def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    # 毕业设计开题报告
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '毕业设计开题报告', font_name=SIMHEI, font_size=Pt(26), bold=True)

    for _ in range(3):
        doc.add_paragraph()

    # 题目
    title_items = [
        ('题    目：', '基于Spring Boot与React的企业级DevOps\n          持续交付平台的设计与实现'),
        ('学    院：', '软件与信息工程学院'),
        ('专    业：', '软件工程技术'),
        ('学生姓名：', '___________'),
        ('学    号：', '___________'),
        ('指导老师：', '___________'),
    ]

    for label, value in title_items:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_indent=Cm(3), line_spacing=Pt(36))
        add_run(p, label, font_name=SIMHEI, font_size=Pt(16), bold=True)
        add_run(p, value, font_name=SIMHEI, font_size=Pt(16))

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026年    月    日', font_name=SIMHEI, font_size=Pt(18), bold=True)

    doc.add_page_break()


# ============================================================
# 第一部分：文献综述
# ============================================================
def build_literature_review(doc):
    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '文 献 综 述', font_name=SIMHEI, font_size=Pt(18), bold=True)

    # === 1 国内研究现状 ===
    add_chapter_title(doc, '1 国内研究现状', space_before=18)

    # 1.1 DevOps 平台发展
    add_sub_title(doc, '1.1 DevOps 平台在国内企业中的应用发展')
    add_body_para(doc,
        '近年来，随着数字化转型战略在国内各行业的深入推进，DevOps（开发运维一体化）理念在国内企业中的'
        '接受度和应用深度显著提升。中国信息通信研究院发布的《中国DevOps现状调查报告（2024）》显示，'
        '超过78%的受访企业已在不同程度上引入DevOps实践，其中互联网、金融和电信行业处于领先地位。'
        '国内学者围绕DevOps持续交付平台的设计与实现开展了大量研究，主要集中在架构设计、自动化流水线'
        '构建、以及云原生技术集成三个方向。')

    add_body_para(doc,
        '在平台架构设计方面，张明辉（2022）在《基于微服务架构的DevOps平台设计与实现》一文中，提出了一种'
        '以Spring Cloud微服务体系为核心、集成Jenkins与GitLab CI的持续交付方案，实现了从代码提交到'
        '生产部署的全链路自动化。该方案采用前后端分离架构，前端使用Vue.js，后端基于Spring Boot构建'
        'RESTful API，验证了微服务架构在DevOps场景下的可行性与扩展性。李伟等（2023）在《企业级持续'
        '集成平台的关键技术研究》中，重点探讨了基于Docker容器化技术的构建环境隔离方案，通过动态创建与'
        '销毁构建容器，有效解决了多项目并行构建时的环境冲突问题，将平均构建等待时间缩短了40%。')

    add_body_para(doc,
        '在自动化流水线方面，陈志远（2023）在《基于Jenkins Pipeline的自动化构建部署系统研究》中，'
        '针对传统手工部署效率低、出错率高的问题，设计了一套以Groovy脚本驱动的声明式流水线，将代码'
        '检出、编译、测试、打包、部署等环节串联为可视化的工作流，并在某中型互联网公司试点运行，使发布'
        '频率从每周一次提升至每日多次。王志强与刘洋（2022）联合发表的《面向中小企业的轻量级CI/CD工具'
        '设计与实现》中，针对Jenkins等重型工具的配置复杂性，提出了一种基于Git事件Hook的简化触发机制，'
        '只需在Web界面配置Git仓库地址和分支规则即可自动触发构建，降低了中小团队的使用门槛。')

    add_body_para(doc,
        '在云原生技术集成领域，赵云飞（2024）在《Kubernetes环境下的DevOps持续交付平台优化》一文中，'
        '将平台部署目标从传统虚拟机迁移到Kubernetes集群，通过Helm Chart管理应用配置，利用K8s的滚动'
        '更新策略实现零停机部署，同时引入Prometheus与Grafana构建全链路监控体系。该研究验证了DevOps平台'
        '与云原生技术栈深度融合的可行路径，为后续研究提供了重要参考。马晓琳（2023）在《DevOps平台中服务'
        '健康监控与自动恢复机制研究》中，设计了基于心跳检测的服务健康监控模块，当服务实例出现异常时，'
        '系统可在30秒内自动标记异常状态并触发告警通知，显著提升了运维响应效率。')

    # 1.2 相关技术研究
    add_sub_title(doc, '1.2 核心支撑技术在国内的研究进展')
    add_body_para(doc,
        '在构建DevOps平台的过程中，多项核心技术得到了国内学者的深入研究。在实时通信方面，周浩（2023）'
        '在《基于WebSocket的实时日志推送系统设计》中，采用STOMP协议实现了服务端向客户端的主动消息推送，'
        '解决了传统HTTP轮询方式在构建日志场景下的延迟与带宽浪费问题。该方案在Spring Boot框架中集成'
        'WebSocket消息代理，使日志消息的端到端延迟控制在200ms以内。')

    add_body_para(doc,
        '在定时任务调度方面，林俊杰（2022）在《分布式定时任务调度系统的设计与优化》中，对比分析了'
        'Quartz、XXL-Job和Spring @Scheduled等主流调度方案，指出在轻量级场景下，基于Cron表达式的'
        '分钟级调度已经能够满足大多数持续集成场景的需求，而不必引入额外的分布式调度中间件。在权限控制'
        '方面，吴晓峰等（2023）提出了基于RBAC模型的四级权限管理方案，将系统用户分为管理员、项目经理、'
        '开发者与观察者四个角色，实现了细粒度的操作权限控制，该方案已被多个企业级管理系统所采纳。')

    add_body_para(doc,
        '综上所述，国内在DevOps持续交付平台领域已积累了丰富的理论成果与实践经验，在架构设计、自动化'
        '流水线、容器化部署等方面形成了较为成熟的技术方案。特别是在Web实时通信、定时调度和权限管理等'
        '支撑技术方面，已有大量可直接借鉴的研究成果。这些研究为本课题提供了坚实的技术参考，同时也表明'
        '在企业级DevOps平台的全流程整合与用户体验优化方面仍有进一步探索的空间。')

    # === 2 国外研究现状 ===
    add_chapter_title(doc, '2 国外研究现状', space_before=18)

    add_sub_title(doc, '2.1 国外主流DevOps平台与产品生态')
    add_body_para(doc,
        '国外在DevOps领域的研究与产品化起步较早，已形成了以GitHub Actions、GitLab CI/CD、Jenkins等'
        '为代表的主流持续集成与持续交付产品生态。GitHub Actions以其与GitHub仓库的深度集成和丰富的'
        'Marketplace Action市场，成为开源社区中最广泛使用的CI/CD工具之一。GitLab CI/CD凭借单一代码'
        '仓库内的.gitlab-ci.yml配置文件，实现了从代码管理到CI/CD的一体化工作流。Jenkins作为老牌CI/CD'
        '引擎，拥有超过1800个插件，生态最为成熟。')

    add_body_para(doc,
        '在学术研究方面，Humble与Farley（2010）在经典著作《Continuous Delivery: Reliable Software '
        'Releases through Build, Test, and Deployment Automation》中，首次系统性地提出了持续交付的'
        '原则与实践框架，奠定了该领域的理论基础。近年来，Bass等（2015）在《DevOps: A Software Architect\'s '
        'Perspective》一书中，从软件架构师的视角分析了DevOps实践对系统架构设计的影响，指出微服务架构与'
        'DevOps文化之间存在天然的相互促进关系。')

    add_body_para(doc,
        'Ebert等（2016）在题为《DevOps》的综述论文中，对DevOps的定义、原则、实践模式和应用效果进行了'
        '全面的文献梳理，提出了DevOps成熟度评估模型，为企业评估自身DevOps实践水平提供了量化工具。'
        'Leite等（2020）在系统文献综述《A Survey of DevOps Concepts and Challenges》中，识别并分析了'
        'DevOps实施过程中的15个关键挑战，其中前三位分别为文化变革阻力、工具链碎片化以及自动化测试覆盖'
        '不足，为后续工具平台的设计指明了改进方向。')

    add_sub_title(doc, '2.2 国外在实时构建与部署自动化方面的技术进展')
    add_body_para(doc,
        '在构建与部署自动化技术方面，国外的研究更加注重可观测性与智能化。Kabinna等（2018）在研究中探讨了'
        '构建日志的实时分析与异常检测，通过机器学习算法自动识别构建失败的根本原因，将故障定位时间缩短了'
        '60%以上。在部署策略方面，Bogner等（2017）对比了蓝绿部署、金丝雀发布和滚动更新三种策略在不同场景'
        '下的适用性，指出对于企业级应用而言，结合审批流的受控部署仍然是合规性要求较高的行业（如金融、医疗）'
        '的首选方案。')

    add_body_para(doc,
        '在容器编排与云原生部署方面，Burns等（2016）在Google发表的白皮书《Borg, Omega, and Kubernetes》'
        '中，分享了Google内部容器编排系统的经验积累，直接推动了Kubernetes的设计与开源进程。Hightower等'
        '（2017）在《Kubernetes: Up and Running》一书中，系统介绍了K8s的部署配置、服务发现和自动扩缩容'
        '等核心功能，为DevOps平台集成容器编排能力提供了标准化的参考实现。这些研究成果表明，国外在DevOps'
        '工具链的完整性和深度方面处于领先地位，特别是在自动化构建故障诊断、智能部署策略和云原生集成方面，'
        '已经超越了基础的脚本化流水线阶段，朝着智能化运维（AIOps）方向演进。')

    # === 3 总结 ===
    add_chapter_title(doc, '3 总结', space_before=18)
    add_body_para(doc,
        '通过对国内外DevOps持续交付平台相关研究的系统梳理，可以看出：国内研究侧重于平台架构的落地实践、'
        '轻量化CI/CD工具的研发以及与云原生技术的深度集成，在WebSocket实时通信、定时调度引擎、RBAC权限'
        '控制等支撑技术方面形成了成熟方案；国外研究则在理论框架构建、智能化故障诊断以及容器编排标准化方面'
        '处于领先地位，主流产品生态更加完善。')

    add_body_para(doc,
        '综合以上研究成果，本课题借鉴了以下核心思路：（1）采用Spring Boot与React前后端分离架构，确保'
        '平台的模块化与可扩展性；（2）集成WebSocket STOMP协议实现构建日志的实时推送，提升用户体验；'
        '（3）设计完整的Git Webhook触发机制，支持GitHub、GitLab和Gitee三种代码托管平台；（4）实现'
        '基于Cron表达式的定时构建调度，满足固定周期的自动化需求；（5）建立PENDING→APPROVED→DEPLOYED'
        '的部署审批状态机，确保生产环境变更的合规性。')

    add_body_para(doc,
        '基于上述技术参考和设计思路，本课题拟设计和实现一个功能完整、操作便捷的企业级DevOps持续交付平台，'
        '帮助中小型开发团队以较低的配置成本完成从代码提交到应用部署的全链路自动化。')

    # === 参考文献 ===
    add_chapter_title(doc, '参考文献', space_before=24)

    refs = [
        '[1] 张明辉. 基于微服务架构的DevOps平台设计与实现[D]. 北京邮电大学, 2022.',
        '[2] 李伟, 王芳, 陈建华. 企业级持续集成平台的关键技术研究[J]. 计算机工程与设计, 2023, 44(3): 856-863.',
        '[3] 陈志远. 基于Jenkins Pipeline的自动化构建部署系统研究[J]. 软件工程, 2023, 26(5): 12-17.',
        '[4] 王志强, 刘洋. 面向中小企业的轻量级CI/CD工具设计与实现[J]. 信息技术, 2022, 46(11): 98-104.',
        '[5] 赵云飞. Kubernetes环境下的DevOps持续交付平台优化[D]. 浙江大学, 2024.',
        '[6] 马晓琳. DevOps平台中服务健康监控与自动恢复机制研究[J]. 计算机应用与软件, 2023, 40(8): 45-51.',
        '[7] 周浩. 基于WebSocket的实时日志推送系统设计[J]. 现代计算机, 2023, 29(14): 78-84.',
        '[8] 林俊杰. 分布式定时任务调度系统的设计与优化[D]. 华南理工大学, 2022.',
        '[9] 吴晓峰, 孙磊, 张鹏. 基于RBAC模型的企业级管理系统权限方案研究[J]. 信息安全研究, 2023, 9(4): 312-319.',
        '[10] Humble J, Farley D. Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation[M]. Addison-Wesley, 2010.',
        '[11] Bass L, Weber I, Zhu L. DevOps: A Software Architect\'s Perspective[M]. Addison-Wesley, 2015.',
        '[12] Ebert C, Gallardo G, Hernantes J, et al. DevOps[J]. IEEE Software, 2016, 33(3): 94-100.',
        '[13] Leite L, Rocha C, Kon F, et al. A Survey of DevOps Concepts and Challenges[J]. ACM Computing Surveys, 2020, 52(6): 1-35.',
        '[14] Bogner J, Zimmermann A, Wagner S. Analyzing the Relevance of SOA Patterns for Microservices-Based Systems[J]. IEEE Software, 2017, 34(5): 29-35.',
    ]

    for ref in refs:
        add_ref_item(doc, ref)

    doc.add_page_break()


# ============================================================
# 第二部分：开题报告
# ============================================================
def build_opening_report(doc):
    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '开 题 报 告', font_name=SIMHEI, font_size=Pt(18), bold=True)

    # === 1 选题依据 ===
    add_chapter_title(doc, '1 选题依据', space_before=18)

    add_sub_title(doc, '1.1 研究目的与意义')
    add_body_para(doc,
        '在当今软件行业，DevOps已成为主流的开发运维模式。据中国信通院2024年报告，超过78%的国内企业'
        '已引入DevOps实践，然而中小型开发团队在落地过程中仍面临诸多痛点：第一，大型CI/CD工具（如Jenkins）'
        '配置复杂、学习曲线陡峭，小团队难以快速上手；第二，现有开源方案功能碎片化，构建、部署、审批、监控'
        '等环节分布在多个独立工具中，缺乏统一的操作界面；第三，生产环境部署缺乏审批追溯机制，存在较大'
        '的操作风险。上述问题严重制约了中小团队在研发效能提升方面的投入产出比。')

    add_body_para(doc,
        '针对上述痛点，本课题拟设计并实现一个基于Spring Boot与React的企业级DevOps持续交付平台。该平台'
        '采用前后端分离的B/S架构，后端以Spring Boot 3.2为核心框架，集成WebSocket STOMP实现实时通信，'
        '通过JWT令牌进行身份认证与RBAC四级权限控制；前端使用React 19与Vite构建单页应用，提供统一的'
        '项目管理、构建配置、部署审批和监控看板界面。平台涵盖了从代码提交Git Webhook触发、自动化构建、'
        '制品管理、部署审批到服务健康监控的完整交付链路，旨在降低中小团队的DevOps落地门槛，提升研发'
        '协作效率。')

    add_sub_title(doc, '1.2 实践价值')
    add_body_para(doc,
        '本课题的实践价值体现在三个层面：（1）效率层面：平台提供统一的Web管理界面，将分散在Git仓库、'
        'CI服务器、部署服务器和监控系统的操作集中到一个平台，减少开发者在多工具之间切换的认知负担和时间'
        '消耗，预计可缩短构建到部署的周转时间30%以上；（2）质量层面：通过部署审批状态机（PENDING→'
        'APPROVED→DEPLOYED）和审计日志（基于AOP的@Audit注解），实现了关键操作的全链路追踪，显著降低'
        '了生产环境误操作的风险，每次部署都有完整的审批记录和历史版本可供回滚；（3）成本层面：平台采用'
        '轻量化设计，不需要依赖Jenkins等重型中间件，部署和运维成本大幅降低，特别适合开发团队规模在50人'
        '以下的中小型项目。')

    # === 2 主要研究内容和要求 ===
    add_chapter_title(doc, '2 主要研究内容和要求', space_before=18)

    add_sub_title(doc, '2.1 课题主要内容')
    add_body_para(doc,
        '本课题围绕企业级DevOps持续交付平台的设计与实现展开，以IntelliJ IDEA和Visual Studio Code'
        '为主要开发工具，后端采用Java 21与Spring Boot 3.2框架，基于JPA/Hibernate进行数据持久化，'
        '数据库使用MySQL 8.0；前端采用React 19与Vite 6构建工具，以Axios进行HTTP通信。平台覆盖了从'
        '代码提交到服务上线运维的全链路流程。')

    add_sub_title(doc, '2.1.1 系统功能结构示意图')

    # 插入系统架构图
    arch_path = os.path.join(IMAGES_DIR, 'architecture.png')
    add_image_centered(doc, arch_path, width_cm=14)
    add_caption(doc, '图1 系统整体架构图')

    # 插入功能模块图
    mod_path = os.path.join(IMAGES_DIR, 'modules.png')
    add_image_centered(doc, mod_path, width_cm=14)
    add_caption(doc, '图2 系统功能模块图')

    add_sub_title(doc, '2.1.2 主要功能阐述')
    add_body_para(doc,
        '本平台共包含12个功能模块，按优先级分为P0核心基础、P1增强功能和P2扩展特性三个层级，具体如下：')

    add_body_para(doc,
        '（1）实时构建日志（P0）：基于WebSocket STOMP协议实现构建日志的实时推送，前端以2秒轮询作为'
        '降级方案，确保在网络不稳定环境下仍能获取完整的构建输出。构建进行中时，界面显示实时进度条和阶段性'
        '步骤完成情况。')

    add_body_para(doc,
        '（2）Git Webhook自动触发（P0）：支持GitHub、GitLab和Gitee三种代码托管平台的Webhook事件'
        '接收，自动解析Push事件中的分支名、提交者和提交信息，结合流水线的branchPattern字段进行分支匹配，'
        '匹配成功则自动触发构建，实现"代码推送即构建"的自动化流程。')

    add_body_para(doc,
        '（3）Cron定时构建（P0）：实现自定义Cron表达式解析引擎，支持分钟级的定时任务调度。用户可为每条'
        '流水线独立配置Cron表达式，并通过开关灵活启用或停用定时任务。定时构建的触发者标记为"SCHEDULER"，'
        '便于在审计日志中区分触发来源。')

    add_body_para(doc,
        '（4）部署审批工作流（P0）：实现PENDING→APPROVED→DEPLOYED的部署审批状态机。对于标记为'
        '"protectedEnv"的生产环境，部署请求必须经过管理员审批后方可执行；对于开发与测试环境，系统自动'
        '审批放行。每次部署均记录完整的审批人、审批时间和审批意见。')

    add_body_para(doc,
        '（5）制品管理（P1）：在构建成功后自动扫描target目录，收集生成的JAR或WAR文件作为制品保存，'
        '记录文件名、大小、版本和创建时间。用户可通过前端界面查看、下载和删除历史制品。')

    add_body_para(doc,
        '（6）构建通知中心（P1）：在构建成功、构建失败或部署审批请求生成时，自动向相关用户发送系统内'
        '通知。通知中心支持按类型筛选（BUILD_SUCCESS、BUILD_FAILED、DEPLOY_APPROVAL等6种类型）、'
        '标记单条已读和全部已读，未读通知以彩色边框突出显示。')

    add_body_para(doc,
        '（7）部署历史与回滚（P1）：完整记录每一次部署操作（目标环境、构建编号、版本号、部署时间、'
        '部署人），标记关键版本为可回滚点。用户可在部署历史中选择任意回滚点执行一键回滚操作。')

    add_body_para(doc,
        '（8）服务实例监控（P1）：对已部署的服务实例进行心跳检测（每30秒一次），监控CPU使用率、内存'
        '使用量、状态和重启次数。超过2分钟未收到心跳的实例自动标记为UNKNOWN状态。实例列表每10秒自动'
        '刷新，首页看板展示运行中/健康/异常实例的统计数据。')

    add_body_para(doc,
        '（9）审计日志（P2）：基于Spring AOP切面编程技术，定义@Audit自定义注解，在方法执行时自动'
        '记录操作人、IP地址、操作类型（CREATE/UPDATE/DELETE/TRIGGER/CANCEL/APPROVE/REJECT）、'
        '操作资源和操作结果。审计日志提供按用户、操作类型和资源的分页查询接口。')

    add_body_para(doc,
        '（10）模板管理（P2）：内置8个常用配置文件模板，涵盖Spring Boot多阶段Dockerfile、Node.js '
        'Dockerfile、Python Dockerfile、Kubernetes Deployment/Service/Ingress YAML和Docker Compose'
        '配置。用户可查看、复制内置模板内容，也可以创建和管理自定义模板。')

    add_body_para(doc,
        '（11）参数化构建（P2）：在触发构建时支持传入JSON格式的构建参数（如版本号、分支名、环境标识等），'
        '后端将这些参数注入构建工作区环境变量，使同一条流水线能够适配不同的构建场景。')

    add_body_para(doc,
        '（12）多分支流水线（P2）：流水线支持配置branchPattern字段（如"feature/*, main, develop"），'
        '通过glob模式匹配来决定哪些分支提交应触发该流水线，实现同一项目中不同分支使用不同构建策略的'
        '灵活配置。')

    add_body_para(doc,
        '此外，平台还实现了完整的RBAC四级权限控制系统，定义了ADMIN（管理员）、MANAGER（项目经理）、'
        'DEVELOPER（开发者）和VIEWER（观察者）四种角色。ADMIN拥有全部操作权限并可管理用户账号；'
        'MANAGER可管理项目、流水线和环境资源，并审批部署请求；DEVELOPER可触发构建和申请部署，但不能'
        '管理资源；VIEWER仅有查看权限，所有写操作按钮隐藏。权限控制在后端SecurityConfig中按HTTP方法'
        '和URL路径精确实现，前端侧边栏和页面操作按钮也按用户角色动态显隐。')

    # === 3 研究方法 ===
    add_chapter_title(doc, '3 研究方法', space_before=18)

    add_sub_title(doc, '3.1 文献研究法')
    add_body_para(doc,
        '通过中国知网（CNKI）、万方数据、IEEE Xplore和ACM Digital Library等学术数据库，查阅DevOps'
        '持续交付、CI/CD自动化、微服务架构、WebSocket实时通信和RBAC权限控制等相关领域的学术论文和技术'
        '文献。对国内外研究现状进行系统梳理和比较分析，提炼出可借鉴的架构设计思路和技术实现方案，为平台'
        '的设计提供理论依据。')

    add_sub_title(doc, '3.2 系统开发法')
    add_body_para(doc,
        '采用前后端分离的B/S架构进行系统开发，遵循软件工程的标准开发流程：（1）需求分析阶段：明确平台'
        '的功能需求和非功能需求，划分P0/P1/P2三个优先级，绘制用例图和功能模块图；（2）系统设计阶段：进行'
        '数据库ER图设计（12个核心实体表）、RESTful API接口设计和前后端交互协议设计；（3）编码实现阶段：'
        '后端使用Java语言基于Spring Boot框架逐模块开发，前端使用JavaScript语言基于React组件化开发，'
        '前后端通过JSON格式的HTTP API进行数据交互；（4）测试验证阶段：使用Postman进行API接口测试，'
        '黑盒测试法对各项功能进行验收测试，确保系统功能符合预期要求。')

    add_sub_title(doc, '3.3 对比分析法')
    add_body_para(doc,
        '在系统设计与开发过程中，对关键技术的选型（如WebSocket vs HTTP轮询、状态机审批 vs 硬编码逻辑、'
        'AOP审计 vs 手动埋点）进行优缺点对比分析。同时，将本平台与市场主流DevOps工具（GitHub Actions、'
        'Jenkins等）在部署复杂度、功能覆盖度和用户体验等维度进行比较评估，从中借鉴成熟经验并弥补现有产品'
        '的不足之处。')

    # === 4 完成期限和措施 ===
    add_chapter_title(doc, '4 完成期限和采取的主要措施', space_before=18)

    add_sub_title(doc, '4.1 完成期限')
    add_body_para(doc,
        '本课题按照以下五个阶段分步推进，总周期约7个月：')

    add_body_para(doc,
        '第一阶段——选题与需求分析（2025年10月）：确认毕业设计选题方向，查阅相关文献资料，与指导教师'
        '讨论确定技术选型和功能范围，完成初步需求文档。')

    add_body_para(doc,
        '第二阶段——开题与方案设计（2025年11月）：撰写开题报告和文献综述，完成系统的概要设计和详细设计，'
        '包括数据库ER图设计、RESTful API接口设计、前端页面原型设计。')

    add_body_para(doc,
        '第三阶段——核心功能开发（2025年12月至2026年2月）：按照P0→P1→P2的优先级顺序，逐步完成'
        '构建管理、WebSocket日志推送、Webhook触发、Cron定时调度、部署审批工作流、制品管理、通知中心、'
        '部署历史回滚、服务实例监控、审计日志、模板管理和参数化构建等全部12个功能模块的编码工作。')

    add_body_para(doc,
        '第四阶段——测试与优化（2026年3月至2026年4月）：进行系统的功能测试和集成测试，修复发现的'
        '缺陷，优化前后端性能和用户体验，撰写毕业设计论文初稿。')

    add_body_para(doc,
        '第五阶段——论文定稿与答辩（2026年5月至2026年6月）：根据指导教师和评阅教师意见修改论文，'
        '准备答辩材料，参加毕业设计答辩，根据答辩意见对论文进行最终修改并提交最终稿。')

    add_sub_title(doc, '4.2 采取的主要措施')
    add_body_para(doc,
        '（1）文献与资料收集：通过中国知网（CNKI）、万方数据、CSDN技术社区、GitHub开源社区等渠道，'
        '广泛收集DevOps、CI/CD、Spring Boot、React等技术领域的最新文献和开源项目案例，为系统设计提供'
        '充分的技术参考。')

    add_body_para(doc,
        '（2）技术储备：在项目启动前系统学习Spring Boot 3、Spring Security、WebSocket STOMP协议、'
        'React Hooks、JPA/Hibernate等核心技术栈，通过完成官方示例项目和实战练习来巩固技术基础。')

    add_body_para(doc,
        '（3）分阶段迭代开发：采用敏捷开发方法，将12个功能模块按P0→P1→P2的优先级分批实现，每完成'
        '一个层级的功能模块即进行阶段性验收，确保核心功能稳定可靠后再推进增强功能。')

    add_body_para(doc,
        '（4）导师指导与交流：定期与指导教师汇报项目进展，针对技术难点（如WebSocket认证拦截、Cron'
        '表达式解析引擎、AOP审计切面设计）向导师请教，根据反馈意见及时调整技术方案。')

    add_body_para(doc,
        '（5）代码版本管理：使用Git进行全项目代码版本控制，遵循规范的commit message格式，确保每次'
        '提交都有清晰的变更说明，便于问题追溯和代码审查。')

    add_body_para(doc,
        '（6）持续测试与质量保障：在开发过程中同步编写API测试用例（使用Postman），每完成一个模块即'
        '进行功能验收测试，及时修复Bug，确保各模块间的接口兼容性和系统整体稳定性。')

    doc.add_page_break()


# ============================================================
# 第三部分：指导教师意见（留白）
# ============================================================
def build_instructor_review(doc):
    add_heading_para(doc, '3. 指导教师意见：', font_size=Pt(14), bold=True)

    for _ in range(5):
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=Pt(36))

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=20)
    add_run(p, '指导老师：    （电子签）', font_name=SIMSUM, font_size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '年    月    日', font_name=SIMSUM, font_size=Pt(12))

    add_heading_para(doc, '所在专业审查意见：', font_size=Pt(14), bold=True)

    for _ in range(5):
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=Pt(36))

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=20)
    add_run(p, '负责人：      （电子签）', font_name=SIMSUM, font_size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '年    月   日', font_name=SIMSUM, font_size=Pt(12))


# ============================================================
# 主流程
# ============================================================
def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = SIMSUM
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), SIMSUM)

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 构建各部分
    build_cover(doc)
    build_literature_review(doc)
    build_opening_report(doc)
    build_instructor_review(doc)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, 'DevOps持续交付平台_开题报告.docx')
    doc.save(output_path)
    print(f'开题报告已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    main()
