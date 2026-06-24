#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成实训报告 - 基于DevOps持续交付平台项目"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ===== 配置 =====
STUDENT_ID = '2023XXXXXXXX'      # 学号 - 请修改
STUDENT_NAME = 'XXX'             # 姓名 - 请修改
CLASS_NAME = '软件工程技术2301班'  # 班级
TEACHER = '李晨子'               # 授课教师
SCORE = '82'                     # 实训成绩

OUTPUT_DIR = r'E:\Dissertation\2\devops-platform\outputs'
DOCX_SRC = r'C:\Users\Dearkoma\Desktop\report_temp.docx'

# 定义中文引号
LQ = '\u201c'  # "
RQ = '\u201d'  # "

# ===== 加载模板 =====
doc = Document(DOCX_SRC)

# ===== 修改 Table 0 - 学生信息表 =====
info_table = doc.tables[0]
info_table.cell(0, 1).text = STUDENT_ID
info_table.cell(1, 1).text = STUDENT_NAME
info_table.cell(2, 1).text = '软件工程技术'
info_table.cell(3, 1).text = CLASS_NAME
info_table.cell(4, 1).text = TEACHER

# ===== 修改 Table 1 - 主内容表 =====
main_table = doc.tables[1]

# Row 1: 实训任务和实训要点
tasks_text = (
    '1. 完成DevOps持续交付平台的系统功能模块识别与划分\n'
    '2. 绘制系统整体功能结构图，展示各模块间的关系\n'
    '3. 描述平台核心功能的用例场景与操作流程\n'
    '4. 分析各功能模块的技术实现方案'
)
main_table.cell(1, 1).text = tasks_text

main_table.cell(1, 3).text = (
    '1. 掌握前后端分离架构设计方法\n'
    '2. 理解CI/CD持续集成交付流程\n'
    '3. 熟练使用UML用例图描述系统功能\n'
    '4. 学会功能模块划分与接口设计'
)

# Row 2: 学生姓名
main_table.cell(2, 1).text = STUDENT_NAME

# Row 3: 成绩
main_table.cell(3, 2).text = SCORE

# ===== Row 4: 实训步骤及内容标题 =====
main_table.cell(4, 1).merge(main_table.cell(4, 4))
main_table.cell(4, 1).text = '基于 Spring Boot + React 的 DevOps 持续交付平台'

# ===== 添加详细内容行 =====

def add_section(table, num, title):
    """添加章节标题行"""
    row = table.add_row()
    row.cells[0].text = num
    row.cells[1].merge(row.cells[4])
    row.cells[1].text = title

def add_content(table, text):
    """添加内容行"""
    row = table.add_row()
    row.cells[0].text = ''
    row.cells[1].merge(row.cells[4])
    row.cells[1].text = text

# ============================
# 第一部分：功能模块识别
# ============================
add_section(main_table, '一', '功能模块识别')

add_content(main_table, (
    f'本实训项目以{LQ}DevOps持续交付平台{RQ}为实训对象，该平台是一个面向企业级软件研发团队的自动化构建、'
    '测试与部署管理系统。通过对平台的全面分析，共识别出12个功能模块，按重要性分为核心模块（P0）、'
    '增强模块（P1）和辅助模块（P2）三个层级。'
))

# P0模块
add_section(main_table, '1', 'P0 核心模块（4项）')

add_content(main_table, (
    '（1）WebSocket实时日志推送：通过STOMP协议实现构建日志的实时推送，前端实时展示构建进度；\n'
    '（2）Git Webhook自动触发构建：支持GitHub/GitLab/Gitee三种平台，代码推送后自动触发CI流程；\n'
    '（3）定时构建模块：基于Cron表达式，支持自定义时间周期的自动化构建任务调度；\n'
    f'（4）部署审批工作流：实现{LQ}申请->审批/驳回->部署->完成{RQ}的完整部署流程，含状态机流转。'
))

# P1模块
add_section(main_table, '2', 'P1 增强模块（4项）')

add_content(main_table, (
    '（1）制品管理：自动扫描构建产物（target/*.jar），支持版本化存储与下载；\n'
    '（2）构建通知中心：支持6种通知类型（构建成功/失败/部署审批/部署成功/失败/系统通知），实时推送与已读标记；\n'
    '（3）部署历史与回滚：记录每次部署的完整快照，支持一键回滚至历史版本；\n'
    '（4）服务实例监控：通过30秒心跳检测机制，实时监控服务实例的运行状态、CPU/内存使用率。'
))

# P2模块
add_section(main_table, '3', 'P2 辅助模块（4项）')

add_content(main_table, (
    '（1）审计日志：基于Spring AOP切面编程，自动记录用户的所有操作行为（增/删/改/触发/审批等），含操作人、IP、时间、结果；\n'
    '（2）模板管理：内置8个CI/CD模板（Dockerfile、K8s Deployment/Service/Ingress、Docker Compose等），支持自定义模板；\n'
    '（3）参数化构建：构建时可动态传入自定义参数（JSON格式），实现灵活的构建配置；\n'
    '（4）多分支流水线：通过glob模式匹配（如feature/*），将不同的Git分支绑定到不同的CI/CD流水线。'
))

# ============================
# 第二部分：功能结构图绘制
# ============================
add_section(main_table, '二', '功能结构图绘制')

add_content(main_table, (
    '以下为DevOps持续交付平台的系统整体架构图，展示四层架构及其子系统组成。\n'
    '（见图1：系统整体架构图）'
))

# 插入架构图
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    fig, ax = plt.subplots(1, 1, figsize=(12, 9), dpi=120)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')

    purple_bg, purple_stroke = '#EEEDFE', '#534AB7'
    blue_bg, blue_stroke = '#E6F1FB', '#185FA5'
    teal_bg, teal_stroke = '#E1F5EE', '#0F6E56'
    coral_bg, coral_stroke = '#FAECE7', '#993C1D'

    def draw_box(ax, x, y, w, h, bg, stroke, text, subtext=None, fs=9, fs_sub=7):
        box = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.1,rounding_size=8',
                             facecolor=bg, edgecolor=stroke, linewidth=0.8, zorder=2)
        ax.add_patch(box)
        ax.text(x + w/2, y + h - 4, text, ha='center', va='top', fontsize=fs,
                fontweight='bold', color='#222', zorder=3)
        if subtext:
            ax.text(x + w/2, y + 3, subtext, ha='center', va='bottom',
                    fontsize=fs_sub, color=stroke, zorder=3)

    def draw_module(ax, x, y, w, h, bg, stroke, title, sub=None):
        box = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.05,rounding_size=5',
                             facecolor=bg, edgecolor=stroke, linewidth=0.5, zorder=2)
        ax.add_patch(box)
        ax.text(x + w/2, y + h - 3.5, title, ha='center', va='top', fontsize=7.5,
                fontweight='bold', color='#222', zorder=3)
        if sub:
            ax.text(x + w/2, y + 1.5, sub, ha='center', va='bottom',
                    fontsize=6.5, color='#666', zorder=3)

    # Title
    ax.text(6, 8.6, 'DevOps 持续交付平台 系统整体架构图', ha='center',
            fontsize=14, fontweight='bold', color='#222')

    # Layer 1: User
    draw_box(ax, 0.5, 7.2, 11, 0.9, purple_bg, purple_stroke,
             '用户层 User Layer',
             '管理员(ADMIN) / 项目经理(MANAGER) / 开发人员(DEVELOPER) / 访客(VIEWER)',
             fs=10)
    ax.text(2, 7.35, '浏览器', ha='center', fontsize=7, color=purple_stroke)
    ax.text(5.5, 7.35, '移动端', ha='center', fontsize=7, color=purple_stroke)
    ax.text(9.5, 7.35, '通知(Web/邮件/钉钉)', ha='center', fontsize=7, color=purple_stroke)

    # Arrow 1->2
    arr = FancyArrowPatch((6, 7.2), (6, 6.5), arrowstyle='->,head_width=6,head_length=4',
                          color=purple_stroke, lw=1, zorder=1)
    ax.add_patch(arr)
    ax.text(6.4, 6.85, 'HTTPS / WebSocket', fontsize=7, color=purple_stroke)

    # Layer 2: Frontend
    draw_box(ax, 0.5, 5.4, 11, 1.1, blue_bg, blue_stroke,
             '前端层 Frontend Layer  React 19 + Vite 6', fs=10)
    fmods = [('监控看板', 'Dashboard'), ('项目管理', 'Projects'),
             ('构建记录', 'Builds'), ('部署管理', 'Deployments'), ('更多', '+8 Pages')]
    for i, (t, s) in enumerate(fmods):
        draw_module(ax, 0.8 + i*2.15, 5.5, 1.95, 0.75, '#B5D4F4', blue_stroke, t, s)

    # Arrow 2->3
    arr = FancyArrowPatch((6, 5.4), (6, 4.7), arrowstyle='->,head_width=6,head_length=4',
                          color=blue_stroke, lw=1, zorder=1)
    ax.add_patch(arr)
    ax.text(6.4, 5.05, 'REST API  JWT Auth', fontsize=7, color=blue_stroke)

    # Layer 3: Backend
    draw_box(ax, 0.5, 3.2, 11, 1.5, teal_bg, teal_stroke,
             '后端层 Backend Layer  Spring Boot 3.2.1 + Java 21', fs=10)
    bmods = [
        (0.7, 3.35, '构建服务', 'BuildService'),
        (0.7, 3.78, '部署服务', 'DeployService'),
        (2.3, 3.35, 'Webhook控制器', 'GitHub/GitLab/Gitee'),
        (2.3, 3.78, '定时调度器', 'Scheduler'),
        (3.9, 3.35, '安全认证', 'JWT + RBAC'),
        (3.9, 3.78, '通知服务', 'Notification'),
        (5.5, 3.35, '审计日志', 'AOP切面'),
        (5.5, 3.78, '实例监控', '30s心跳'),
        (7.1, 3.35, '制品管理', 'Artifact'),
        (7.1, 3.78, '模板管理', 'Template'),
    ]
    for x, y, t, s in bmods:
        draw_module(ax, x, y, 1.4, 0.36, '#9FE1CB', teal_stroke, t, s)

    # WebSocket badge
    box = FancyBboxPatch((7.5, 3.1), 3.5, 0.22,
                         boxstyle='round,pad=0.05,rounding_size=4',
                         facecolor=teal_stroke, edgecolor='none', zorder=3)
    ax.add_patch(box)
    ax.text(9.25, 3.21, 'WebSocket 实时日志推送 (STOMP)', ha='center',
            fontsize=7, color='white', zorder=4)

    # Arrow 3->4
    arr = FancyArrowPatch((3, 3.2), (2, 2.4), arrowstyle='->,head_width=6,head_length=4',
                          color=teal_stroke, lw=1, zorder=1)
    ax.add_patch(arr)
    arr = FancyArrowPatch((9, 3.2), (10, 2.4), arrowstyle='->,head_width=6,head_length=4',
                          color=teal_stroke, lw=1, zorder=1)
    ax.add_patch(arr)

    # Layer 4: Data
    draw_box(ax, 0.5, 1.0, 11, 1.0, coral_bg, coral_stroke,
             '数据层 Data Layer', fs=10)
    dmods = [('MySQL 8.0', 'JPA/Hibernate'), ('制品库', 'Artifacts'),
             ('审计日志存储', 'AuditLogs'), ('文件存储', 'FileSystem'),
             ('内存缓存', 'LogCache')]
    for i, (t, s) in enumerate(dmods):
        draw_module(ax, 0.8 + i*2.1, 1.15, 1.8, 0.45, '#F5C4B3', coral_stroke, t, s)

    # External Git
    box = FancyBboxPatch((9.2, 7.25), 2.3, 0.65,
                         boxstyle='round,pad=0.1,rounding_size=6',
                         facecolor='#FBEAF0', edgecolor='#993556',
                         linewidth=0.8, linestyle='dashed', zorder=2)
    ax.add_patch(box)
    ax.text(10.35, 7.7, '外部 Git 平台', ha='center', fontsize=8,
            fontweight='bold', color='#4B1528', zorder=3)
    ax.text(10.35, 7.4, 'GitHub  GitLab  Gitee', ha='center',
            fontsize=6.5, color='#72243E', zorder=3)

    # Git -> Webhook arrow (curved)
    from matplotlib.patches import ConnectionPatch
    con = ConnectionPatch(xyA=(10.35, 7.25), xyB=(2.8, 3.7),
                          coordsA='data', coordsB='data',
                          arrowstyle='->,head_width=6,head_length=4',
                          color='#993556', lw=0.8, linestyle='dashed',
                          connectionstyle='arc3,rad=-0.3', zorder=1)
    ax.add_artist(con)
    ax.text(7.2, 5.7, 'Git Webhook 推送', fontsize=7, color='#993556')

    # Tech stack footer
    ax.text(6, 0.3, '技术栈: React 19  Vite 6  Spring Boot 3.2.1  Java 21  MySQL 8.0  JPA/Hibernate  WebSocket STOMP  Maven 3.9',
            ha='center', fontsize=7, color='#999')

    png_arch_path = os.path.join(OUTPUT_DIR, 'arch_diagram_report.png')
    plt.tight_layout(pad=0.5)
    plt.savefig(png_arch_path, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    print(f'Architecture diagram saved: {png_arch_path}')

    # Insert into document
    row_img = main_table.add_row()
    row_img.cells[0].text = ''
    row_img.cells[1].merge(row_img.cells[4])
    p = row_img.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_arch_path, width=Inches(5.8))
    print('Diagram inserted into document')

except Exception as e:
    print(f'Warning: Could not generate diagram: {e}')
    import traceback
    traceback.print_exc()
    row_img = main_table.add_row()
    row_img.cells[0].text = ''
    row_img.cells[1].merge(row_img.cells[4])
    row_img.cells[1].text = '（系统架构图请参见 outputs/系统整体架构图.svg）'

# 架构图说明
add_content(main_table, (
    '图1：系统整体架构图\n\n'
    '平台采用经典的四层B/S架构设计：\n\n'
    f'用户层：支持管理员(ADMIN)、项目经理(MANAGER)、开发人员(DEVELOPER)、访客(VIEWER)四级角色，'
    '通过JWT令牌实现无状态认证与RBAC细粒度权限控制。不同角色拥有差异化的功能访问权限，'
    '管理员拥有全功能权限，访客仅可查看不可操作。\n\n'
    '前端层：基于React 19框架与Vite 6构建工具，实现12个功能页面（监控看板、项目管理、构建记录、'
    '定时任务、制品管理、部署管理、环境管理、通知中心、服务实例、模板管理、审计日志、用户管理），'
    '通过React Router 7实现SPA单页路由，通过Axios与后端REST API通信，'
    'WebSocket/轮询双模式接收实时构建日志。\n\n'
    '后端层：基于Spring Boot 3.2.1框架（Java 21），核心服务包括：'
    'BuildService（构建执行引擎，支持Maven/npm/Docker编译，8种步骤类型，真实/模拟双执行模式）、'
    'DeploymentService（部署审批状态机，PENDING->APPROVED/REJECTED->DEPLOYED/FAILED）、'
    'SchedulerService（Cron定时调度，每分钟检查自定义Cron表达式）、'
    'NotificationService（6种通知类型，实时消息推送）、'
    'InstanceMonitorService（服务实例30秒心跳检测）、'
    'SecurityConfig（Spring Security + JWT四级RBAC权限控制）。\n\n'
    '数据层：MySQL 8.0关系型数据库，通过JPA/Hibernate ORM框架管理12张核心业务表'
    '（users, projects, pipelines, builds, environments, artifacts, '
    'deployment_requests, deployment_history, notifications, service_instances, audit_logs, templates），'
    '支持自动DDL更新（ddl-auto: update）。'
))

# 前端模块结构图
add_content(main_table, (
    '前端功能模块结构（12个页面）：\n\n'
    '| 监控看板(Dashboard) | 项目管理(Projects) | 构建记录(Builds) | 定时任务(Scheduler) |\n'
    '| 制品管理(Artifacts) | 部署管理(Deployments) | 环境管理(Environments) | 通知中心(Notifications) |\n'
    '| 服务实例(Instances) | 模板管理(Templates) | 审计日志(AuditLog) | 用户管理(Users) |\n\n'
    f'其中：{LQ}审计日志{RQ}和{LQ}用户管理{RQ}为管理员(ADMIN)专属模块；'
    f'{LQ}项目管理{RQ}、{LQ}环境管理{RQ}、{LQ}模板管理{RQ}需管理员或项目经理(MANAGER)权限；'
    f'{LQ}构建记录{RQ}、{LQ}部署管理{RQ}、{LQ}制品管理{RQ}对开发者(DEVELOPER)及以上开放；'
    f'{LQ}监控看板{RQ}对访客(VIEWER)及以上开放。'
))

# ============================
# 第三部分：核心功能用例描述
# ============================
add_section(main_table, '三', '核心功能用例描述')

# 用例1
add_section(main_table, '1', '【用例UC-01】用户认证与权限管理')

add_content(main_table, (
    '参与角色：所有用户（管理员/项目经理/开发者/访客）\n'
    '前置条件：数据库中存在有效的用户账号\n\n'
    '基本路径：\n'
    f'  步骤1：用户访问登录页面，输入用户名和密码，点击{LQ}登录{RQ}按钮；\n'
    '  步骤2：后端验证用户名密码（BCrypt加密匹配），通过后生成JWT令牌（HS256签名、24小时有效期、含角色信息），返回给前端；\n'
    '  步骤3：前端将令牌存储在localStorage中，后续所有API请求在Authorization头中携带Bearer Token；\n'
    '  步骤4：JwtAuthFilter拦截每个请求，解析Token并注入Spring Security上下文，完成无状态认证；\n'
    '  步骤5：SecurityConfig根据用户角色（ADMIN/MANAGER/DEVELOPER/VIEWER），按HTTP方法+URL路径精确匹配，控制功能访问权限。\n\n'
    '异常路径：\n'
    '  用户名或密码错误 -> 返回401 + 错误提示；\n'
    '  Token过期(24小时) -> 返回401，前端自动跳转到登录页；\n'
    f'  无权限访问 -> 返回403 {LQ}权限不足{RQ}。\n\n'
    '后置条件：用户成功登录后，根据角色显示对应的侧边栏导航菜单和功能按钮。'
))

# 用例2
add_section(main_table, '2', '【用例UC-02】手动触发CI构建')

add_content(main_table, (
    '参与角色：管理员、项目经理、开发者\n'
    '前置条件：已创建项目并配置了Git仓库地址和构建命令\n\n'
    '基本路径：\n'
    f'  步骤1：用户在构建记录页面点击{LQ}触发构建{RQ}按钮，弹出构建参数配置窗口（可选：自定义构建参数JSON、指定分支）；\n'
    '  步骤2：前端调用POST /api/builds/trigger接口，传递projectId、branch、buildParams等参数；\n'
    '  步骤3：BuildService异步执行构建流程（@Async），执行步骤依次为：\n'
    '    a. 工作区准备：检查项目目录，若不存在则git clone，已存在则git fetch + checkout + pull；\n'
    '    b. 流水线执行：逐阶段逐步骤执行Pipeline定义中的每个Step（SHELL/MVN_PACKAGE/NPM_INSTALL/TEST/DOCKER_BUILD等8种类型）；\n'
    '    c. 日志推送：每执行一个步骤，通过WebSocket向/topic/builds/{id}/log推送JSON格式的日志更新（含阶段名、步骤名、状态、输出内容）；\n'
    '    d. 制品收集：构建成功后自动扫描target或dist目录，将生成的jar/war等文件保存到制品库；\n'
    '  步骤4：前端通过STOMP WebSocket订阅构建日志通道，实时渲染构建进度（显示阶段进度条、步骤状态、彩色日志输出），若WebSocket连接失败则降级为HTTP轮询；\n'
    '  步骤5：构建完成（成功/失败/取消）后，触发NotificationService发送通知给相关用户。\n\n'
    '异常路径：\n'
    f'  Git仓库不可达 -> 构建失败，日志显示{LQ}Git clone failed{RQ}；\n'
    '  编译错误 -> 构建失败，日志显示编译错误信息；\n'
    '  构建超时（超过10分钟）-> 自动取消构建。\n\n'
    '后置条件：构建记录写入builds表，制品保存到artifacts表，构建日志保存到文件。'
))

# 用例3
add_section(main_table, '3', '【用例UC-03】部署申请与审批工作流')

add_content(main_table, (
    '参与角色：开发者（申请部署）、管理员/项目经理（审批部署）\n'
    '前置条件：存在构建成功的制品，目标环境已配置\n\n'
    '基本路径：\n'
    '  步骤1：开发者在部署管理页面选择构建产物和目标环境（开发/测试/预发布/生产），填写部署原因，提交部署申请；\n'
    '  步骤2：后端创建DeploymentRequest记录，状态设为PENDING（待审批），若目标环境为受保护环境（如生产环境）则必须经过审批；\n'
    f'  步骤3：有审批权限的用户（ADMIN/MANAGER）在{LQ}审批列表{RQ}中查看待审批的部署申请，可选择{LQ}批准{RQ}或{LQ}驳回{RQ}；\n'
    '  步骤4：批准后：DeploymentRequest状态更新为APPROVED -> 自动执行部署 -> 状态更新为DEPLOYED，同时创建DeploymentHistory记录（含回滚快照）；\n'
    '        驳回后：状态更新为REJECTED，记录驳回原因，通知申请人。\n'
    '  步骤5：部署历史记录完整保存每次部署的版本号、部署URL、操作人、时间戳，支持一键回滚。\n\n'
    '状态流转图（状态机）：\n'
    '  PENDING（待审批）-> APPROVED（已批准）-> DEPLOYED（已部署）\n'
    '  PENDING（待审批）-> REJECTED（已驳回）\n'
    '  DEPLOYED -> ROLLED_BACK（通过回滚操作）\n\n'
    '后置条件：部署记录持久化，通知相关人员，制品下载链接更新。'
))

# 用例4
add_section(main_table, '4', '【用例UC-04】Git Webhook自动触发构建')

add_content(main_table, (
    '参与角色：外部Git平台（GitHub/GitLab/Gitee）\n'
    '前置条件：项目已配置Git仓库URL，CI/CD流水线已关联项目\n\n'
    '基本路径：\n'
    '  步骤1：开发者在IDE中提交代码并推送到Git仓库（如git push origin main）；\n'
    '  步骤2：Git平台检测到Push事件，向平台Webhook端点发送POST请求（含JSON格式的事件载荷，包含仓库信息、分支名、提交者、commit SHA等）；\n'
    '  步骤3：WebhookController接收到Webhook请求后：\n'
    '    a. 解析事件载荷，提取分支名和commit信息；\n'
    '    b. 根据projectId查找对应的Pipeline配置；\n'
    '    c. 检查branchPattern是否匹配当前推送的分支（如feature/*匹配feature/login分支）；\n'
    '    d. 匹配成功后自动调用BuildService触发构建（triggerType=PUSH）；\n'
    '  步骤4：后续构建流程同UC-02。\n\n'
    '支持的Webhook格式：\n'
    '  GitHub：x-github-event: push，载荷包含ref、commits、repository等字段；\n'
    '  GitLab：X-Gitlab-Event: Push Hook，载荷包含ref、commits、project等字段；\n'
    '  Gitee：X-Gitee-Event: Push Hook，载荷包含ref、commits、repository等字段。\n\n'
    '后置条件：自动创建的构建记录包含triggerType=PUSH、gitCommit=commit SHA。'
))

# 用例5
add_section(main_table, '5', '【用例UC-05】定时构建调度')

add_content(main_table, (
    '参与角色：系统（自动执行）\n'
    '前置条件：Pipeline已配置cronExpression并启用cronEnabled=true\n\n'
    '基本路径：\n'
    '  步骤1：SchedulerService通过@Scheduled注解每分钟执行一次检查；\n'
    '  步骤2：遍历所有cronEnabled=true的Pipeline，使用自定义Cron解析器解析cronExpression（5段式：分 时 日 月 周）；\n'
    '  步骤3：将解析出的目标触发时间与当前时间匹配，若匹配成功则触发构建；\n'
    '  步骤4：构建执行流程同UC-02（triggerType=SCHEDULE）。\n\n'
    'Cron表达式示例：\n'
    f'  {LQ}0 2 * * *{RQ}     -> 每天凌晨2点构建\n'
    f'  {LQ}*/30 * * * *{RQ}  -> 每30分钟构建一次\n'
    f'  {LQ}0 9 * * 1-5{RQ}  -> 工作日每天早上9点构建\n\n'
    '后置条件：构建记录包含triggerType=SCHEDULE。'
))

# ============================
# 第四部分：实训总结
# ============================
add_section(main_table, '四', '实训总结')

add_content(main_table, (
    f'通过本次实训，我以{LQ}DevOps持续交付平台{RQ}为实训对象，系统性地完成了软件系统的功能模块识别、'
    '功能结构图绘制和核心功能用例描述三个实训任务。\n\n'
    '在功能模块识别环节，我按照软件工程方法论，将平台划分为12个功能模块，并根据业务重要性分为P0（核心）、'
    'P1（增强）、P2（辅助）三个优先级层级，明确了每个模块的职责边界和依赖关系。\n\n'
    '在功能结构图绘制环节，我绘制了系统的四层架构图（用户层->前端层->后端层->数据层）和前端功能模块结构图，'
    '清晰展示了各层组件的技术选型、模块划分和数据流动关系。平台采用Spring Boot + React的前后端分离架构，'
    '通过JWT + RBAC实现四级角色的权限控制，通过WebSocket实现实时通信，通过状态机模式实现部署审批工作流。\n\n'
    '在核心功能用例描述环节，我选取了5个核心用例（用户认证与权限管理、手动触发CI构建、'
    '部署申请与审批工作流、Git Webhook自动触发、定时构建调度）进行详细描述，'
    '涵盖了基本路径、异常路径、前置条件、后置条件等完整用例要素，展示了平台的核心业务逻辑。\n\n'
    '本次实训加深了我对软件系统分析与设计的理解，特别是在大型企业级系统的功能分解、'
    '架构设计和用例建模方面积累了实践经验。同时，通过对Spring Boot、React、WebSocket、'
    'JWT认证、MySQL等主流技术的应用分析，提升了我的技术文档编写能力和系统设计能力。'
))

# ===== 保存文档 =====
output_path = os.path.join(OUTPUT_DIR, '第1次实训报告-学号姓名.docx')
doc.save(output_path)
print(f'实训报告已保存到: {output_path}')
print(f'文件大小: {os.path.getsize(output_path):,} bytes')
print('请修改学号、姓名等信息后提交。')
